from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings

load_dotenv()

urls = [
    "https://lilianweng.github.io/posts/2023-06-23-agent/",
    "https://lilianweng.github.io/posts/2023-03-15-prompt-engineering/",
    "https://lilianweng.github.io/posts/2023-10-25-adv-attack-llm/",
]

docs = [WebBaseLoader(url).load() for url in urls]
docs_list = [item for sublist in docs for item in sublist]

text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
    chunk_size=250, chunk_overlap=0
)
doc_splits = text_splitter.split_documents(docs_list)

vectorstore = Chroma.from_documents(
     documents=doc_splits,
     collection_name="rag-chroma",
     embedding=OpenAIEmbeddings(),
     persist_directory="./.chroma",
 )

retriever = Chroma(
    collection_name="rag-chroma",
    persist_directory="./.chroma",
    embedding_function=OpenAIEmbeddings(),
).as_retriever()


# New helper to process local files uploaded via the dashboard
def process_documents(file_paths: list[str]):
    """
    Load local files (txt, pdf, docx), split them into chunks and add to the existing Chroma collection.
    """
    from langchain.schema import Document as LangDocument

    loaded_docs = []
    for path in file_paths:
        try:
            if path.lower().endswith('.txt'):
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            elif path.lower().endswith('.pdf'):
                # pypdf is used to extract text from PDFs
                from pypdf import PdfReader

                reader = PdfReader(path)
                pages = [p.extract_text() or "" for p in reader.pages]
                text = "\n".join(pages)
            elif path.lower().endswith('.docx') or path.lower().endswith('.doc'):
                # python-docx for docx files
                from docx import Document as DocxDocument

                doc = DocxDocument(path)
                paras = [p.text for p in doc.paragraphs]
                text = "\n".join(paras)
            else:
                # Fallback try to read as text
                with open(path, 'rb') as f:
                    try:
                        text = f.read().decode('utf-8', errors='ignore')
                    except Exception:
                        text = ""

            if text:
                loaded_docs.append(LangDocument(page_content=text, metadata={"source": path}))
        except Exception as e:
            # ignore problematic files but continue
            print(f"Failed to load {path}: {e}")

    if not loaded_docs:
        print("No documents loaded to index.")
        return

    # split and add
    splits = text_splitter.split_documents(loaded_docs)
    try:
        vectorstore.add_documents(splits)
        # persist to disk if supported
        if hasattr(vectorstore, 'persist'):
            try:
                vectorstore.persist()
            except Exception:
                pass
        print(f"Indexed {len(splits)} document chunks into Chroma.")
    except Exception as e:
        print(f"Failed to add documents to vectorstore: {e}")
